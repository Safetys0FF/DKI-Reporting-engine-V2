#!/usr/bin/env python3
"""Adapter for Mission Debrief report generator integration."""

from __future__ import annotations

import importlib
import json
import logging
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

# Ensure we can attempt a clean import from preferred locations
sys.modules.pop("report_generator", None)

_REPORT_GENERATOR_SEARCH_PATHS = [
    Path(__file__).resolve().parents[1] / "report generator",
    Path(__file__).resolve().parents[3] / "The War Room" / "Processors" / "report generator",
]

ReportGenerator = None
_GENERATOR_AVAILABLE = False
_GENERATOR_IMPORT_ERROR: Optional[Exception] = None

for search_path in _REPORT_GENERATOR_SEARCH_PATHS:
    if not search_path.exists():
        continue
    path_str = str(search_path)
    if path_str not in sys.path:
        sys.path.insert(0, path_str)
    try:
        module = importlib.import_module("report_generator")
        ReportGenerator = getattr(module, "ReportGenerator")  # type: ignore[attr-defined]
        _GENERATOR_AVAILABLE = True
        _GENERATOR_IMPORT_ERROR = None
        break
    except Exception as exc:  # pragma: no cover - diagnostic import branch
        _GENERATOR_IMPORT_ERROR = exc
        continue

logger = logging.getLogger(__name__)


def _module_available(module_name: str) -> bool:
    try:
        importlib.import_module(module_name)
        return True
    except Exception:  # pragma: no cover - dependency probe
        return False


class ReportGeneratorAdapter:
    """Expose report generation and export helpers to Mission Debrief."""

    def __init__(self) -> None:
        self.logger = logging.getLogger(__name__)
        self.generator: Optional[Any] = None
        self.import_error: Optional[Exception] = _GENERATOR_IMPORT_ERROR
        self.generator_available = _GENERATOR_AVAILABLE and ReportGenerator is not None
        self.output_dir = self._determine_output_dir()
        self.have_docx = _module_available("docx")
        self.have_reportlab = _module_available("reportlab")

        if ReportGenerator is not None:
            try:
                self.generator = ReportGenerator(output_dir=str(self.output_dir))  # type: ignore[call-arg]
                self.output_dir = Path(getattr(self.generator, "output_dir", self.output_dir))
            except TypeError:
                try:
                    self.generator = ReportGenerator()  # type: ignore[call-arg]
                    self.output_dir = Path(getattr(self.generator, "output_dir", self.output_dir))
                except Exception as exc:
                    self.logger.exception("Report generator initialisation failed")
                    self.generator = None
                    self.import_error = exc
            except Exception as exc:  # pragma: no cover - defensive
                self.logger.exception("Report generator initialisation failed")
                self.generator = None
                self.import_error = exc

        self.output_dir = Path(self.output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.available = True

    # ------------------------------------------------------------------
    # Capability helpers
    # ------------------------------------------------------------------
    def is_available(self) -> bool:
        return self.available

    def capability_status(self) -> Dict[str, Any]:
        return {
            "available": self.is_available(),
            "generator_loaded": self.generator is not None,
            "docx_supported": self.have_docx,
            "pdf_supported": self.have_reportlab,
            "import_error": str(self.import_error) if self.import_error else None,
            "output_dir": str(self.output_dir),
        }

    # ------------------------------------------------------------------
    # Generation
    # ------------------------------------------------------------------
    def generate(self, section_data: Any, report_type: str = "Investigative") -> Dict[str, Any]:
        sections = self._normalise_sections(section_data)
        timestamp = datetime.now()
        case_id = self._resolve_case_id(section_data) or f"CASE-{timestamp:%Y%m%d%H%M%S}"

        cover_page = self._build_cover_page(case_id, report_type, timestamp)
        table_of_contents = self._build_table_of_contents(sections)
        disclosure_page = self._build_disclosure_page(case_id, report_type, timestamp)

        report_payload: Dict[str, Any] = {
            "cover_page": cover_page,
            "table_of_contents": table_of_contents,
            "sections": sections,
            "disclosure_page": disclosure_page,
            "metadata": {
                "case_id": case_id,
                "report_type": report_type,
                "generated_timestamp": timestamp.isoformat(),
                "total_sections": len(sections),
                "artifact": "Final Report",
            },
        }

        self._append_raw_generation(section_data, case_id, report_payload)

        report_text = self._compose_report_text(report_payload)
        report_payload["report_text"] = report_text
        report_path = report_payload.get("report_path")
        if report_path:
            try:
                Path(report_path).write_text(report_text, encoding="utf-8")
            except Exception as exc:
                self.logger.debug("Unable to update report text at %s: %s", report_path, exc)
        else:
            report_payload["report_path"] = self._write_text_report(report_payload, case_id, timestamp)

        report_payload.setdefault("generated_at", timestamp.strftime("%Y-%m-%d %H:%M:%S"))
        report_payload.setdefault("status", "ok")
        report_payload["case_id"] = case_id
        report_payload["output_files"] = [report_payload.get("report_path")]
        return report_payload

    def _append_raw_generation(self, section_data: Any, case_id: str, report_payload: Dict[str, Any]) -> None:
        if not self.generator:
            return
        evidence = None
        if isinstance(section_data, dict):
            evidence = section_data.get("evidence")
        try:
            sections_map = {
                entry["section_id"]: entry.get("content", "")
                for entry in report_payload.get("sections", [])
                if isinstance(entry, dict) and entry.get("section_id")
            }
            result = self.generator.generate_full_report(
                evidence=evidence if isinstance(evidence, dict) else None,
                sections=sections_map,
                case_id=case_id,
            )
        except Exception as exc:  # pragma: no cover - defensive
            self.logger.warning("Fallback text generation engaged: %s", exc)
            report_payload.setdefault("status", "warning")
            report_payload.setdefault("errors", []).append(str(exc))
            return
        if not isinstance(result, dict):
            return
        report_path = result.get("report_path")
        if report_path:
            report_payload["report_path"] = report_path
        status = result.get("status")
        if status == "error":
            report_payload.setdefault("errors", []).append(result.get("error"))
            report_payload["status"] = "warning"
        elif status:
            report_payload["status"] = status
        generated_at = result.get("generated_at")
        if generated_at:
            report_payload["generated_at"] = generated_at

    # ------------------------------------------------------------------
    # Export
    # ------------------------------------------------------------------
    def export(self, payload: Dict[str, Any], export_path: str, export_format: str = "PDF") -> Dict[str, Any]:
        export_path_obj = Path(export_path)
        export_path_obj.parent.mkdir(parents=True, exist_ok=True)
        fmt = (export_format or export_path_obj.suffix.replace(".", "")).upper() if (export_format or export_path_obj.suffix) else "PDF"
        if fmt in {"PDF"} and self.have_reportlab:
            self._export_pdf(payload, export_path_obj)
        elif fmt in {"DOCX", "MS_WORD"} and self.have_docx:
            self._export_docx(payload, export_path_obj)
        else:
            if fmt == "PDF" and not self.have_reportlab:
                self.logger.warning("reportlab not available; exporting plain text instead of PDF")
            if fmt in {"DOCX", "MS_WORD"} and not self.have_docx:
                self.logger.warning("python-docx not available; exporting plain text instead of DOCX")
            self._export_text(payload, export_path_obj)
            fmt = "TEXT"
        return {
            "status": "ok",
            "output_path": str(export_path_obj),
            "format": fmt,
        }

    def _export_text(self, payload: Dict[str, Any], export_path: Path) -> None:
        export_path.write_text(self._compose_report_text(payload), encoding="utf-8")

    def _export_pdf(self, payload: Dict[str, Any], export_path: Path) -> None:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        doc = SimpleDocTemplate(
            str(export_path),
            pagesize=letter,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )
        styles = getSampleStyleSheet()
        story: List[Any] = []

        def add_block(title: str, text: str) -> None:
            if not text and not title:
                return
            if title:
                story.append(Paragraph(title, styles["Heading2"]))
                story.append(Spacer(1, 0.15 * inch))
            if text:
                paragraph = text.replace("\n", "<br/>")
                story.append(Paragraph(paragraph, styles["BodyText"]))
                story.append(Spacer(1, 0.2 * inch))

        add_block("Cover Page", payload.get("cover_page", {}).get("content", ""))
        add_block("Table of Contents", payload.get("table_of_contents", {}).get("content", ""))
        for section in payload.get("sections", []) or []:
            if not isinstance(section, dict):
                continue
            title = section.get("title") or section.get("section_id") or "Section"
            add_block(title, section.get("content", ""))
        add_block("Disclosure", payload.get("disclosure_page", {}).get("content", ""))
        if not story:
            add_block("Report", "No content available.")
        doc.build(story)

    def _export_docx(self, payload: Dict[str, Any], export_path: Path) -> None:
        from docx import Document

        doc = Document()

        def add_section(title: str, text: str, add_break: bool = True) -> None:
            if title:
                doc.add_heading(title, level=1)
            if text:
                for paragraph in text.splitlines():
                    doc.add_paragraph(paragraph)
            else:
                doc.add_paragraph("")
            if add_break:
                doc.add_page_break()

        sections = payload.get("sections", []) or []
        add_section("Cover Page", payload.get("cover_page", {}).get("content", ""))
        add_section("Table of Contents", payload.get("table_of_contents", {}).get("content", ""))
        for index, section in enumerate(sections):
            if not isinstance(section, dict):
                continue
            title = section.get("title") or section.get("section_id") or f"Section {index + 1}"
            add_section(title, section.get("content", ""))
        add_section("Disclosure", payload.get("disclosure_page", {}).get("content", ""), add_break=False)
        doc.save(str(export_path))

    # ------------------------------------------------------------------
    # Snapshot helpers
    # ------------------------------------------------------------------
    def save_snapshot(self, payload: Dict[str, Any], output_dir: str) -> str:
        directory = Path(output_dir)
        directory.mkdir(parents=True, exist_ok=True)
        snapshot_path = directory / f"report_snapshot_{datetime.now():%Y%m%d_%H%M%S}.json"
        snapshot_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
        return str(snapshot_path)

    def load_snapshot(self, snapshot_path: str) -> Dict[str, Any]:
        path = Path(snapshot_path)
        return json.loads(path.read_text(encoding="utf-8"))

    # ------------------------------------------------------------------
    # Normalisation helpers
    # ------------------------------------------------------------------
    def _normalise_sections(self, section_data: Any) -> List[Dict[str, Any]]:
        sections: List[Dict[str, Any]] = []
        if isinstance(section_data, dict):
            iterator = section_data.items()
        elif isinstance(section_data, list):
            iterator = []
            for idx, entry in enumerate(section_data, start=1):
                iterator.append((entry.get("section_id", f"section_{idx}") if isinstance(entry, dict) else f"section_{idx}", entry))
        else:
            return sections
        for key, value in iterator:
            normalised = self._normalise_single_section(key, value)
            if normalised:
                sections.append(normalised)
        return sections

    def _normalise_single_section(self, key: Any, value: Any) -> Optional[Dict[str, Any]]:
        if isinstance(value, dict):
            section_id = str(value.get("section_id") or key)
            title = self._resolve_section_title(value, section_id)
            content = self._resolve_section_content(value)
            metadata: Dict[str, Any] = {}
            meta_block = value.get("metadata")
            if isinstance(meta_block, dict):
                metadata.update(meta_block)
            case_data = value.get("case_data")
            if isinstance(case_data, dict):
                metadata.setdefault("case_id", case_data.get("case_id"))
            return {
                "section_id": section_id,
                "title": title,
                "content": content,
                "metadata": metadata,
            }
        if value is None:
            return None
        section_id = str(key)
        return {
            "section_id": section_id,
            "title": section_id.replace("_", " ").title(),
            "content": str(value),
            "metadata": {},
        }

    def _resolve_section_title(self, value: Dict[str, Any], fallback: str) -> str:
        for field in ("title", "section_title", "name", "label"):
            text = value.get(field)
            if isinstance(text, str) and text.strip():
                return text.strip()
        return fallback.replace("_", " ").title()

    def _resolve_section_content(self, value: Dict[str, Any]) -> str:
        for field in ("content", "narrative", "summary", "body", "text"):
            data = value.get(field)
            if isinstance(data, str) and data.strip():
                return data.strip()
        structured = value.get("structured_data")
        if isinstance(structured, dict):
            lines: List[str] = []
            for key, item in structured.items():
                if isinstance(item, (str, int, float)):
                    lines.append(f"{key}: {item}")
                elif isinstance(item, dict):
                    for sub_key, sub_value in item.items():
                        if isinstance(sub_value, (str, int, float)):
                            lines.append(f"{key}.{sub_key}: {sub_value}")
            if lines:
                return "\n".join(lines)
        return ""

    def _resolve_case_id(self, section_data: Any) -> Optional[str]:
        if isinstance(section_data, dict):
            direct = section_data.get("case_id")
            if direct:
                return str(direct)
            for value in section_data.values():
                case_id = self._extract_case_id_from_entry(value)
                if case_id:
                    return case_id
        elif isinstance(section_data, list):
            for entry in section_data:
                case_id = self._extract_case_id_from_entry(entry)
                if case_id:
                    return case_id
        return None

    def _extract_case_id_from_entry(self, entry: Any) -> Optional[str]:
        if not isinstance(entry, dict):
            return None
        candidate = entry.get("case_id")
        if candidate:
            return str(candidate)
        metadata = entry.get("metadata")
        if isinstance(metadata, dict):
            candidate = metadata.get("case_id")
            if candidate:
                return str(candidate)
        case_data = entry.get("case_data")
        if isinstance(case_data, dict):
            candidate = case_data.get("case_id")
            if candidate:
                return str(candidate)
        return None

    def _build_cover_page(self, case_id: str, report_type: str, timestamp: datetime) -> Dict[str, Any]:
        lines = [
            "DKI Services LLC",
            "Mission Debrief",
            "",
            f"Case ID: {case_id}",
            f"Report Type: {report_type}",
            f"Prepared: {timestamp.strftime('%Y-%m-%d %H:%M:%S')}",
        ]
        return {
            "type": "cover_page",
            "content": "\n".join(lines),
            "metadata": {
                "case_id": case_id,
                "report_type": report_type,
                "generated_timestamp": timestamp.isoformat(),
            },
        }

    def _build_table_of_contents(self, sections: List[Dict[str, Any]]) -> Dict[str, Any]:
        lines = ["TABLE OF CONTENTS", ""]
        if sections:
            for index, section in enumerate(sections, start=1):
                title = section.get("title") or section.get("section_id") or f"Section {index}"
                lines.append(f"{index}. {title}")
        else:
            lines.append("No sections available")
        return {
            "type": "table_of_contents",
            "content": "\n".join(lines),
            "metadata": {"total_sections": len(sections)},
        }

    def _build_disclosure_page(self, case_id: str, report_type: str, timestamp: datetime) -> Dict[str, Any]:
        lines = [
            "Disclosure",
            "",
            "This report is intended for authorised personnel within the Central Command network.",
            "Unauthorised distribution or alteration is prohibited.",
            "Findings reflect the information available at the time of preparation.",
            "",
            f"Case ID: {case_id}",
            f"Report Type: {report_type}",
            f"Certified: {timestamp.strftime('%Y-%m-%d %H:%M:%S')}",
        ]
        return {
            "type": "disclosure_page",
            "content": "\n".join(lines),
            "metadata": {
                "case_id": case_id,
                "report_type": report_type,
                "certification_timestamp": timestamp.isoformat(),
            },
        }

    def _compose_report_text(self, payload: Dict[str, Any]) -> str:
        segments: List[str] = []
        for key in ("cover_page", "table_of_contents"):
            section = payload.get(key)
            if isinstance(section, dict):
                content = section.get("content")
                if content:
                    segments.append(content.strip())
        for section in payload.get("sections", []) or []:
            if not isinstance(section, dict):
                continue
            title = section.get("title") or section.get("section_id")
            content = section.get("content", "")
            block_lines = []
            if title:
                block_lines.append(f"## {title}")
            if content:
                block_lines.append(content.strip())
            if block_lines:
                segments.append("\n".join(block_lines))
        disclosure = payload.get("disclosure_page")
        if isinstance(disclosure, dict):
            content = disclosure.get("content")
            if content:
                segments.append(content.strip())
        return "\n\n".join(segment for segment in segments if segment) or "No report content available."

    def _write_text_report(self, payload: Dict[str, Any], case_id: str, timestamp: datetime) -> str:
        safe_case = self._sanitise_case_id(case_id)
        filename = f"{safe_case}_FinalReport_{timestamp:%Y%m%d_%H%M%S}.txt"
        destination = self.output_dir / filename
        destination.write_text(payload.get("report_text", ""), encoding="utf-8")
        return str(destination)

    def _sanitise_case_id(self, case_id: str) -> str:
        return "".join(ch if ch.isalnum() or ch in {"-", "_"} else "_" for ch in case_id)

    def _determine_output_dir(self) -> Path:
        primary = Path("F:/The Central Command/Generated Reports")
        try:
            primary.mkdir(parents=True, exist_ok=True)
            return primary
        except Exception:
            fallback = Path(__file__).resolve().parents[1] / "productions" / "reports"
            fallback.mkdir(parents=True, exist_ok=True)
            return fallback


__all__ = ["ReportGeneratorAdapter"]
