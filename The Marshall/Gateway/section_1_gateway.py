# Section 1 Renderer (Gateway)

from typing import Dict, Any, Tuple, List, Optional

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False


class Section1Renderer:
    """
    Rendering-only engine for Section 1 – Investigation Objectives / Case Info.
    - Fonts/Styles locked
    - Whitelist (drift guard)
    - Placeholder policy (italicized at output)
    - Produces a section-only artifact + manifest
    - Implements 3x cyclical fallback validation across intake, notes, evidence, and prior section
    """

    SECTION_KEY = "section_1"
    TITLE = "SECTION 1 – INVESTIGATION OBJECTIVES / CASE INFO"

    WHITELIST = [
        "client_name", "client_address", "client_phone", "contract_date",
        "investigation_goals",
        "subject_primary", "subject_secondary", "subject_tertiary",
        "subject_employers", "subject_employer_address",
        "agency_name", "agency_license",
        "assigned_investigator", "investigator_license",
        "location_of_investigation",
    ]

    LAYOUT = [
        ("CLIENT INFORMATION", [
            ("Client Name", "client_name"),
            ("Client Address", "client_address"),
            ("Client Phone", "client_phone"),
            ("Date of Contract", "contract_date"),
        ]),
        ("GOALS OF INVESTIGATION", [
            ("Goals of Investigation", "investigation_goals"),
        ]),
        ("SUBJECTS OF INVESTIGATION", [
            ("Primary Subject", "subject_primary"),
            ("Secondary Subject", "subject_secondary"),
            ("Tertiary Subject", "subject_tertiary"),
            ("Employer(s)", "subject_employers"),
            ("Employer Address", "subject_employer_address"),
        ]),
        ("AGENCY AND LICENSE", [
            ("Agency Name", "agency_name"),
            ("Agency License", "agency_license"),
        ]),
        ("ASSIGNED INVESTIGATOR", [
            ("Investigator Name", "assigned_investigator"),
            ("Investigator License", "investigator_license"),
        ]),
        ("LOCATION OF INVESTIGATION", [
            ("Location", "location_of_investigation"),
        ]),
    ]

    PLACEHOLDERS = {
        "unknown": "Unknown",
        "unconfirmed": "Unconfirmed at this time",
        "suppressed": "Due to the nature of this case this portion was not performed or was not necessary",
    }
    BANNED_TOKENS = { "", " ", "N/A", "NA", "TBD", "[REDACTED]", "REDACTED" }

    def __init__(self):
        pass

    def _normalize(self, v: Optional[Any]) -> Optional[str]:
        if v is None:
            return None
        s = str(v).strip()
        return s if s else None

    def _placeholder_for(self, v: Optional[str]) -> Tuple[str, bool]:
        if v is None or (v.upper() in self.BANNED_TOKENS):
            return (self.PLACEHOLDERS["unknown"], True)
        lower = v.lower()
        if lower in {"pending", "not verified"}:
            return (self.PLACEHOLDERS["unconfirmed"], True)
        return (v, False)

    def _fallback_check(self, key, zones):
        for _ in range(3):
            for zone in ["intake", "notes", "evidence", "prior_section"]:
                val = zones.get(zone, {}).get(key)
                if val:
                    return val
        return None

    def render_model(self, section_payload: Dict[str, Any], case_sources: Dict[str, Any] = None) -> Dict[str, Any]:
        data: Dict[str, Any] = section_payload or {}
        drift_bounced: Dict[str, Any] = {
            k: v for k, v in data.items() if k not in self.WHITELIST
        }

        rendered_blocks: List[Dict[str, Any]] = []
        placeholders_used: Dict[str, str] = {}

        rendered_blocks.append({
            "type": "title",
            "text": self.TITLE,
            "style": {"font": "Times New Roman", "size_pt": 16, "bold": True,
                      "all_caps": True, "align": "center", "spacing": 1.15}
        })

        for subheader, fields in self.LAYOUT:
            rendered_blocks.append({
                "type": "header",
                "text": subheader,
                "style": {"font": "Times New Roman", "size_pt": 14, "bold": True,
                          "underline": True, "all_caps": True, "align": "left", "spacing": 1.15,
                          "space_before_lines": 2, "space_after_lines": 0}
            })
            for label, key in fields:
                if key not in self.WHITELIST:
                    continue
                val = data.get(key)
                if not val and case_sources:
                    val = self._fallback_check(key, case_sources)
                raw = self._normalize(val)
                text, is_ph = self._placeholder_for(raw)
                if is_ph:
                    placeholders_used[key] = text
                rendered_blocks.append({
                    "type": "field",
                    "label": label,
                    "value": text,
                    "style": {
                        "font": "Times New Roman",
                        "size_pt": 12,
                        "label_bold": True,
                        "value_italic": is_ph,
                        "align": "left",
                        "spacing": 1.15
                    }
                })

        manifest = {
            "section_key": self.SECTION_KEY,
            "title": self.TITLE,
            "fields_rendered": [k for _, grp in self.LAYOUT for _, k in grp],
            "placeholders_used": placeholders_used,
            "drift_bounced": drift_bounced
        }

        return {
            "render_tree": rendered_blocks,
            "manifest": manifest,
            "handoff": "gateway"
        }

    def render_docx(self, section_payload: Dict[str, Any], out_path: str) -> Dict[str, Any]:
        model = self.render_model(section_payload)
        if not HAVE_DOCX:
            return {"artifact": None, "note": "python-docx not available", **model}

        doc = Document()

        t = doc.add_paragraph()
        tr = t.add_run(self.TITLE.upper())
        tr.bold = True
        tr.font.name = "Times New Roman"
        tr.font.size = Pt(16)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for block in model["render_tree"][1:]:
            if block["type"] == "header":
                p = doc.add_paragraph()
                r = p.add_run(block["text"].upper())
                r.bold = True
                r.underline = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)
                doc.add_paragraph()
            elif block["type"] == "field":
                p = doc.add_paragraph()
                r1 = p.add_run(f"{block['label']}: ")
                r1.bold = True
                r1.font.name = "Times New Roman"
                r1.font.size = Pt(12)
                r2 = p.add_run(str(block["value"]))
                r2.italic = block["style"]["value_italic"]
                r2.font.name = "Times New Roman"
                r2.font.size = Pt(12)

        doc.save(out_path)
        return {"artifact": out_path, **model}
