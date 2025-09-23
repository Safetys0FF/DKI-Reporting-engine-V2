#!/usr/bin/env python3
"""
Report Generator - Final assembly and export functionality for DKI Engine reports
Handles report compilation, deduplication, formatting, and export to various formats
"""

import os
import logging
from datetime import datetime
from typing import Dict, List, Any, Optional
from pathlib import Path
import json

# Document generation libraries
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image, Table, TableStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.lib import colors
    HAVE_REPORTLAB = True
except ImportError:
    HAVE_REPORTLAB = False

logger = logging.getLogger(__name__)

class ReportGenerator:
    """Comprehensive report generation and export system"""
    
    def __init__(self):
        self.company_info = {
            'name': 'DKI Services LLC',
            'license': '0200812-IA000307',
            'address': 'Tulsa, Oklahoma',
            'phone': '(918) 882-5539',
            'email': 'david@dkiservicesok.com',
            'logo_path': None  # Path to company logo if available
        }
        
        self.investigator_info = {
            'name': 'David Krashin',
            'license': '0163814-C000480',
            'title': 'Licensed Private Investigator',
            'signature_path': None  # Path to signature image if available
        }
        
        self.report_templates = {
            'Investigative': self._get_investigative_template(),
            'Surveillance': self._get_surveillance_template(),
            'Hybrid': self._get_hybrid_template()
        }
        
        logger.info("Report Generator initialized")
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check and log available export capabilities"""
        deps = {
            'DOCX Export': HAVE_DOCX,
            'PDF Export': HAVE_REPORTLAB
        }
        
        for dep, available in deps.items():
            status = "Available" if available else "Missing"
            logger.info(f"{dep}: {status}")
    
    def generate_full_report(self, section_data: Dict[str, Any], report_type: str) -> Dict[str, Any]:
        """Generate complete report from section data"""
        
        logger.info(f"Generating full {report_type} report")
        
        # Get report template
        template = self.report_templates.get(report_type, self.report_templates['Investigative'])
        
        # Compile sections in proper order
        compiled_sections = self._compile_sections(section_data, template['section_order'])
        
        # Deduplicate content
        deduplicated_content = self._deduplicate_content(compiled_sections)
        
        # Generate cover page
        cover_page = self._generate_cover_page(section_data, report_type)
        
        # Generate table of contents
        toc = self._generate_table_of_contents(deduplicated_content)
        
        # Generate disclosure page
        disclosure_page = self._generate_disclosure_page(section_data, report_type)
        
        # Assemble final report
        final_report = {
            'cover_page': cover_page,
            'table_of_contents': toc,
            'sections': deduplicated_content,
            'disclosure_page': disclosure_page,
            'metadata': {
                'report_type': report_type,
                'generated_timestamp': datetime.now().isoformat(),
                'total_sections': len(deduplicated_content),
                'case_id': self._extract_case_id(section_data),
                'client_name': self._extract_client_name(section_data)
            }
        }
        
        logger.info(f"Full report generated with {len(deduplicated_content)} sections")
        
        return final_report
    
    def _get_investigative_template(self) -> Dict[str, Any]:
        """Template for investigative reports"""
        return {
            'section_order': [
                'section_1',   # Investigation Objectives
                'section_2',   # Investigation Requirements
                'section_3',   # Investigation Details
                'section_4',   # Review of Details
                'section_5',   # Review of Supporting Documents
                'section_9',   # Certification & Disclaimers
                'section_7',   # Conclusion
                'section_8',   # Investigation Evidence Review
                'section_6'    # Billing Summary (near end)
            ],
            'title_format': 'Investigation Report',
            'subtitle_format': 'Comprehensive Investigation Analysis'
        }
    
    def _get_surveillance_template(self) -> Dict[str, Any]:
        """Template for surveillance reports"""
        return {
            'section_order': [
                'section_1',   # Surveillance Objectives
                'section_2',   # Pre-Surveillance Planning
                'section_3',   # Investigation Details
                'section_4',   # Review of Surveillance Sessions
                'section_5',   # Review of Supporting Documents
                'section_9',   # Certification & Disclaimers
                'section_7',   # Conclusion
                'section_8',   # Investigation Evidence Review
                'section_6'    # Billing Summary
            ],
            'title_format': 'Surveillance Report',
            'subtitle_format': 'Professional Surveillance Documentation'
        }
    
    def _get_hybrid_template(self) -> Dict[str, Any]:
        """Template for hybrid reports"""
        return {
            'section_order': [
                'section_1',   # Investigation Objectives
                'section_2',   # Preliminary Case Review
                'section_3',   # Investigative Details
                'section_4',   # Review of Surveillance Sessions
                'section_5',   # Review of Supporting Documents
                'section_9',   # Certification & Disclaimers
                'section_7',   # Conclusion
                'section_8',   # Investigation Evidence Review
                'section_6'    # Billing Summary
            ],
            'title_format': 'Comprehensive Investigation Report',
            'subtitle_format': 'Combined Investigation and Surveillance Analysis'
        }
    
    def _compile_sections(self, section_data: Dict[str, Any], section_order: List[str]) -> List[Dict[str, Any]]:
        """Compile sections in the specified order"""
        
        compiled = []
        
        for section_id in section_order:
            # Look for section data in various possible formats
            section_content = None
            
            # Check direct section ID
            if section_id in section_data:
                section_content = section_data[section_id]
            
            # Check for section names
            for key, data in section_data.items():
                if isinstance(data, dict) and data.get('section_id') == section_id:
                    section_content = data
                    break
                elif section_id.replace('_', ' ').lower() in key.lower():
                    section_content = data
                    break
            
            if section_content:
                # Ensure content is in proper format
                if isinstance(section_content, dict):
                    compiled.append(section_content)
                elif isinstance(section_content, str):
                    # Convert string content to proper format
                    compiled.append({
                        'section_id': section_id,
                        'content': section_content,
                        'metadata': {'converted_from_string': True}
                    })
            else:
                logger.warning(f"Section {section_id} not found in section data")
        
        return compiled
    
    def _deduplicate_content(self, sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Remove duplicate content across sections"""
        
        deduplicated = []
        seen_content = set()
        
        for section in sections:
            content = section.get('content', '')
            if not content:
                continue
            
            # Split content into paragraphs
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            unique_paragraphs = []
            
            for paragraph in paragraphs:
                # Create a hash of the paragraph for comparison
                paragraph_hash = hash(paragraph.lower().strip())
                
                if paragraph_hash not in seen_content:
                    seen_content.add(paragraph_hash)
                    unique_paragraphs.append(paragraph)
            
            # Reconstruct section with unique content
            if unique_paragraphs:
                section_copy = section.copy()
                section_copy['content'] = '\n\n'.join(unique_paragraphs)
                section_copy['metadata'] = section_copy.get('metadata', {})
                section_copy['metadata']['deduplicated'] = True
                section_copy['metadata']['original_paragraphs'] = len(paragraphs)
                section_copy['metadata']['final_paragraphs'] = len(unique_paragraphs)
                
                deduplicated.append(section_copy)
        
        logger.info(f"Deduplication complete: {len(sections)} -> {len(deduplicated)} sections")
        
        return deduplicated
    
    def _generate_cover_page(self, section_data: Dict[str, Any], report_type: str) -> Dict[str, Any]:
        """Generate report cover page"""
        
        case_id = self._extract_case_id(section_data)
        client_name = self._extract_client_name(section_data)
        date_generated = datetime.now().strftime("%B %d, %Y")
        
        template = self.report_templates[report_type]
        
        # Extract profile data from Section CP manifest (primary source)
        cp = self._find_section_cp(section_data)
        
        # Default values (fallback only)
        inv_name = 'David Krashin'
        inv_title = 'Licensed Private Investigator'
        inv_license = '0163814-C000480'
        co_name = 'DKI Services LLC'
        co_license = '0200812-IA000307'
        co_address = 'Tulsa, Oklahoma'
        
        # Override with user profile data if available
        if cp and isinstance(cp, dict):
            manifest = cp.get('render_data', {}).get('manifest', {}) if 'render_data' in cp else cp.get('manifest', {})
            profile = manifest.get('cover_profile', {})
            
            # Use user profile data, fallback to defaults only if not provided
            inv_name = profile.get('investigator_name') or inv_name
            inv_title = profile.get('investigator_title') or inv_title
            inv_license = profile.get('investigator_license') or inv_license
            co_name = profile.get('agency_name') or co_name
            co_license = profile.get('agency_license') or co_license
            co_address = profile.get('agency_mailing_address') or co_address
            
            logger.debug(f"Cover page using profile data: {inv_name} - {co_name}")
        
        cover_content = f"""
{template['title_format']}

{template['subtitle_format']}

Case ID: {case_id}
Client: {client_name}
Date: {date_generated}

Prepared by:
{inv_name}
{inv_title}
License: {inv_license}

{co_name}
License: {co_license}
{co_address}
"""
        
        return {
            'type': 'cover_page',
            'content': cover_content.strip(),
            'metadata': {
                'case_id': case_id,
                'client_name': client_name,
                'report_type': report_type,
                'date_generated': date_generated
            }
        }
    
    def _generate_table_of_contents(self, sections: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Generate table of contents"""
        
        toc_entries = []
        page_number = 3  # Start after cover page and TOC
        
        for section in sections:
            section_title = self._extract_section_title(section)
            toc_entries.append(f"{section_title} ... {page_number}")
            
            # Estimate pages based on content length
            content_length = len(section.get('content', ''))
            estimated_pages = max(1, content_length // 2000)  # Rough estimate
            page_number += estimated_pages
        
        toc_content = "TABLE OF CONTENTS\n\n" + '\n'.join(toc_entries)
        
        return {
            'type': 'table_of_contents',
            'content': toc_content,
            'metadata': {
                'total_sections': len(sections),
                'estimated_pages': page_number
            }
        }
    
    def _generate_disclosure_page(self, section_data: Dict[str, Any], report_type: str) -> Dict[str, Any]:
        """Generate disclosure and signature page"""
        
        case_id = self._extract_case_id(section_data)
        date_generated = datetime.now().strftime("%B %d, %Y")

        # Extract profile data from Section CP manifest (primary source)
        cp = self._find_section_cp(section_data)
        
        # Default values (fallback only)
        inv_name = 'David Krashin'
        inv_title = 'Licensed Private Investigator'
        inv_license = '0163814-C000480'
        co_name = 'DKI Services LLC'
        co_license = '0200812-IA000307'
        co_phone = '(918) 882-5539'
        co_email = 'david@dkiservicesok.com'
        slogan = ''
        d_logo_path = None
        
        # Override with user profile data if available
        if cp and isinstance(cp, dict):
            manifest = cp.get('render_data', {}).get('manifest', {}) if 'render_data' in cp else cp.get('manifest', {})
            profile = manifest.get('cover_profile', {})
            
            # Use user profile data, fallback to defaults only if not provided
            inv_name = profile.get('investigator_name') or inv_name
            inv_title = profile.get('investigator_title') or inv_title
            inv_license = profile.get('investigator_license') or inv_license
            co_name = profile.get('agency_name') or co_name
            co_license = profile.get('agency_license') or co_license
            co_phone = profile.get('phone') or co_phone
            co_email = profile.get('email') or co_email
            slogan = profile.get('slogan') or slogan
            d_logo_path = profile.get('logo_path') or d_logo_path
            
            logger.debug(f"Disclosure page using profile data: {inv_name} - {co_name}")

        disclosure_content = f"""
DISCLOSURE AND CERTIFICATION

This report has been prepared by {inv_name}, a licensed private investigator 
in the state of Oklahoma (License #{inv_license}). All investigative activities were conducted
in compliance with applicable state statutes and regulations.

The information contained in this report is based on investigation conducted between the dates 
specified herein. All observations, findings, and conclusions are based on available evidence 
and professional analysis. This report is confidential and prepared exclusively for the client.
Unauthorized distribution is prohibited.

Case ID: {case_id}
Report Type: {report_type}
Date of Report: {date_generated}

Disclosures
The documents included with this report include but are not limited to: the client contract, the new client
intake form, the final report, and the comprehensive report on the subjects. The combination of these documents
is to be supplied to the client and handled as sensitive and private information. Please keep the results private
and in a secure location, as sensitive information is included within the documents. Neither I, {inv_name}, nor {co_name}
is responsible if any or all of these documents are mishandled, lost, stolen, or used in an inappropriate manner.

I, {inv_name} of {co_name}, do not portray or identify myself as a lawyer or attorney, nor do I offer any legal advice on
any case. The following statements are none other than the perception and collection of details, facts (to the best of my
knowledge), and records of this specific case. The facts and records found in this report are subject to change without notice.
They are also my thoughts, not facts, that I perceive are likely based upon the events I have witnessed in multiple cases
throughout my career. These statements are not a record of fact and should not be regarded as such. I strongly encourage every
client to confer with a legal representative and adhere to their legal counsel pending the review of this investigation.


_________________________________
{inv_name}
{inv_title}
License: {inv_license}
Date: {date_generated}


_________________________________
{co_name}
License: {co_license}
{('Phone: ' + co_phone) if co_phone else ''}
{('Email: ' + co_email) if co_email else ''}
{('“' + slogan + '”') if slogan else ''}
"""
        
        return {
            'type': 'disclosure_page',
            'content': disclosure_content.strip(),
            'metadata': {
                'case_id': case_id,
                'report_type': report_type,
                'certification_date': date_generated,
                'logo_path': d_logo_path
            }
        }
    
    def export_report(self, report_data: Dict[str, Any], filename: str, format_type: str):
        """Export report to specified format"""
        
        if format_type.lower() == 'pdf':
            self._export_pdf(report_data, filename)
        elif format_type.lower() == 'docx':
            self._export_docx(report_data, filename)
        else:
            raise ValueError(f"Unsupported export format: {format_type}")
        
        logger.info(f"Report exported to {filename}")

    def export_rtf(self, report_data: Dict[str, Any], filename: str):
        """Export a simple RTF version for broad compatibility (Word/Apple)."""
        def esc(s: str) -> str:
            return s.replace('\\', r'\\').replace('{', r'\{').replace('}', r'\}')
        parts = []
        parts.append(r"{\rtf1\ansi\deff0\fs24")
        # Cover
        parts.append(esc(report_data.get('cover_page', {}).get('content', '')) + '\n\n')
        # TOC
        parts.append(esc(report_data.get('table_of_contents', {}).get('content', '')) + '\n\n')
        # Sections
        for section in report_data.get('sections', []) or []:
            title = self._extract_section_title(section)
            parts.append(esc(title) + '\n')
            parts.append(esc(section.get('content', '')) + '\n\n')
        # Disclosure
        parts.append(esc(report_data.get('disclosure_page', {}).get('content', '')) + '\n')
        parts.append('}')
        with open(filename, 'w', encoding='utf-8') as f:
            f.write('\n'.join(parts))
    
    def _export_docx(self, report_data: Dict[str, Any], filename: str):
        """Export report to DOCX format"""
        
        if not HAVE_DOCX:
            raise RuntimeError("python-docx library not available")
        
        doc = Document()
        
        # Set up styles
        self._setup_docx_styles(doc)
        
        # Attempt to set page margins to 0.5"
        try:
            for sec in doc.sections:
                sec.left_margin = Inches(0.5)
                sec.right_margin = Inches(0.5)
                sec.top_margin = Inches(0.5)
                sec.bottom_margin = Inches(0.5)
        except Exception:
            pass
        
        # Add cover page
        self._add_docx_cover_page(doc, report_data['cover_page'])
        doc.add_page_break()
        
        # Add table of contents
        self._add_docx_section(doc, report_data['table_of_contents'])
        doc.add_page_break()
        
        # Add sections
        for section in report_data['sections']:
            self._add_docx_section(doc, section)
            doc.add_page_break()
        
        # Add disclosure page
        self._add_docx_section(doc, report_data['disclosure_page'])
        
        # Save document
        doc.save(filename)
    
    def _export_pdf(self, report_data: Dict[str, Any], filename: str):
        """Export report to PDF format"""
        
        if not HAVE_REPORTLAB:
            raise RuntimeError("reportlab library not available")
        
        doc = SimpleDocTemplate(filename, pagesize=letter, leftMargin=0.5*inch, rightMargin=0.5*inch, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=12,
            alignment=TA_LEFT
        )
        
        # Add cover page (optional logo)
        from pathlib import Path as _Path
        logo_path = (self.company_info or {}).get('logo_path')
        if logo_path and _Path(logo_path).exists():
            try:
                story.append(Image(logo_path, width=2.5*inch))
                story.append(Spacer(1, 12))
            except Exception:
                pass
        cover_lines = report_data['cover_page']['content'].split('\n')
        for line in cover_lines:
            if line.strip():
                story.append(Paragraph(line, title_style))
            else:
                story.append(Spacer(1, 12))
        
        story.append(PageBreak())
        
        # Add table of contents
        toc_lines = report_data['table_of_contents']['content'].split('\n')
        for line in toc_lines:
            if line.strip():
                if 'TABLE OF CONTENTS' in line:
                    story.append(Paragraph(line, title_style))
                else:
                    story.append(Paragraph(line, styles['Normal']))
            else:
                story.append(Spacer(1, 12))
        
        story.append(PageBreak())
        
        # Add sections
        for section in report_data['sections']:
            render_data = section.get('render_data') or {}
            render_tree = render_data.get('render_tree') if isinstance(render_data, dict) else None
            if render_tree:
                buffer_images = []
                def flush_grid_pdf():
                    if not buffer_images:
                        return
                    cells = []
                    row = []
                    for idx, c in enumerate(buffer_images, 1):
                        elems = []
                        if c.get('path') and os.path.exists(c['path']):
                            try:
                                elems.append(Image(c['path'], width=3.25*inch))
                            except Exception:
                                pass
                        if c.get('label'):
                            elems.append(Paragraph(c['label'], styles['Normal']))
                        if c.get('timestamp'):
                            elems.append(Paragraph(c['timestamp'], styles['Normal']))
                        if c.get('address'):
                            elems.append(Paragraph(c['address'], styles['Normal']))
                        if c.get('note'):
                            elems.append(Paragraph(c['note'], styles['Italic']))
                        row.append(elems)
                        if idx % 2 == 0:
                            cells.append(row)
                            row = []
                    if row:
                        while len(row) < 2:
                            row.append("")
                        cells.append(row)
                    t = Table(cells, colWidths=[3.5*inch, 3.5*inch])
                    t.setStyle(TableStyle([
                        ('VALIGN', (0,0), (-1,-1), 'TOP'),
                        ('BOX', (0,0), (-1,-1), 0.25, colors.grey),
                        ('INNERGRID', (0,0), (-1,-1), 0.25, colors.lightgrey),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 12))
                    buffer_images.clear()

                # Walk blocks
                for block in render_tree:
                    btype = block.get('type')
                    if btype == 'title':
                        story.append(Paragraph(block.get('text', ''), header_style))
                    elif btype == 'header':
                        flush_grid_pdf()
                        story.append(Paragraph(block.get('text', ''), header_style))
                    elif btype == 'paragraph':
                        flush_grid_pdf()
                        t = block.get('text', '')
                        if t.strip():
                            story.append(Paragraph(t, styles['Normal']))
                            story.append(Spacer(1, 8))
                    elif btype == 'image':
                        img_path = block.get('path')
                        if img_path and os.path.exists(img_path):
                            buffer_images.append({
                                'path': img_path,
                                'label': block.get('label'),
                                'timestamp': block.get('timestamp'),
                                'address': block.get('address'),
                                'note': block.get('note'),
                            })
                            if len(buffer_images) == 4:
                                flush_grid_pdf()
                flush_grid_pdf()
                story.append(PageBreak())
            else:
                section_title = self._extract_section_title(section)
                story.append(Paragraph(section_title, header_style))
                content_paragraphs = section.get('content', '').split('\n\n')
                for paragraph in content_paragraphs:
                    if paragraph.strip():
                        story.append(Paragraph(paragraph, styles['Normal']))
                        story.append(Spacer(1, 12))
                story.append(PageBreak())
        
        # Add disclosure page (optional logo and signature)
        # Logo at top - use path from disclosure metadata if present (same directive as cover)
        from pathlib import Path as _Path
        logo_path = (report_data.get('disclosure_page', {}).get('metadata', {}) or {}).get('logo_path') or (self.company_info or {}).get('logo_path')
        if logo_path and _Path(logo_path).exists():
            try:
                story.append(Spacer(1, 12))
                story.append(Image(logo_path, width=2.0*inch))
                story.append(Spacer(1, 12))
            except Exception:
                pass
        disclosure_lines = report_data['disclosure_page']['content'].split('\n')
        for line in disclosure_lines:
            if line.strip():
                if 'DISCLOSURE AND CERTIFICATION' in line:
                    story.append(Paragraph(line, title_style))
                else:
                    story.append(Paragraph(line, styles['Normal']))
            else:
                story.append(Spacer(1, 12))
        sig_path = (self.investigator_info or {}).get('signature_path')
        if sig_path and _Path(sig_path).exists():
            try:
                story.append(Spacer(1, 12))
                story.append(Image(sig_path, width=2.0*inch))
            except Exception:
                pass
        
        # Build PDF
        doc.build(story)
    
    def _setup_docx_styles(self, doc):
        """Setup custom styles for DOCX document"""
        
        styles = doc.styles
        
        # Title style
        if 'Custom Title' not in styles:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(18)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(24)
        
        # Header style
        if 'Custom Header' not in styles:
            header_style = styles.add_style('Custom Header', WD_STYLE_TYPE.PARAGRAPH)
            header_style.font.name = 'Times New Roman'
            header_style.font.size = Pt(14)
            header_style.font.bold = True
            header_style.paragraph_format.space_after = Pt(12)
    
    def _add_docx_cover_page(self, doc, cover_page_data: Dict[str, Any]):
        """Add cover page to DOCX document"""
        from pathlib import Path as _Path
        # Optional company logo at top
        logo_path = (self.company_info or {}).get('logo_path')
        if logo_path and _Path(logo_path).exists():
            try:
                doc.add_picture(logo_path, width=Inches(2.5))
                doc.add_paragraph()
            except Exception:
                pass

        lines = cover_page_data['content'].split('\n')
        
        for line in lines:
            if line.strip():
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style = doc.styles['Custom Title']
            else:
                doc.add_paragraph()
    
    def _add_docx_section(self, doc, section_data: Dict[str, Any]):
        """Add section to DOCX document"""

        render_data = section_data.get('render_data') or {}
        render_tree = render_data.get('render_tree') if isinstance(render_data, dict) else None

        if render_tree:
            buffer_images = []
            def flush_grid():
                if not buffer_images:
                    return
                tbl = doc.add_table(rows=0, cols=2)
                # Add rows of two cells
                it = iter(buffer_images)
                for left in it:
                    right = next(it, None)
                    row = tbl.add_row().cells
                    for idx, cell_data in enumerate([left, right]):
                        if not cell_data:
                            continue
                        # Image
                        try:
                            pimg = row[idx].add_paragraph()
                            pimg.alignment = 1  # center
                            r = pimg.add_run()
                            r.add_picture(cell_data['path'], width=Inches(3.25))
                        except Exception:
                            pass
                        # Captions/lines
                        if cell_data.get('label'):
                            row[idx].add_paragraph(cell_data['label'])
                        if cell_data.get('timestamp'):
                            row[idx].add_paragraph(cell_data['timestamp'])
                        if cell_data.get('address'):
                            row[idx].add_paragraph(cell_data['address'])
                        if cell_data.get('note'):
                            n = row[idx].add_paragraph(cell_data['note'])
                            for run in n.runs:
                                run.italic = True
                buffer_images.clear()

            # Walk blocks
            for block in render_tree:
                btype = block.get('type')
                if btype == 'title':
                    p = doc.add_paragraph(block.get('text', ''))
                    p.style = doc.styles['Custom Header']
                elif btype == 'header':
                    flush_grid()
                    p = doc.add_paragraph(block.get('text', ''))
                    p.style = doc.styles['Custom Header']
                elif btype == 'paragraph':
                    flush_grid()
                    t = block.get('text', '')
                    if t.strip():
                        doc.add_paragraph(t)
                elif btype == 'image':
                    img_path = block.get('path')
                    if img_path and os.path.exists(img_path):
                        buffer_images.append({
                            'path': img_path,
                            'label': block.get('label'),
                            'timestamp': block.get('timestamp'),
                            'address': block.get('address'),
                            'note': block.get('note'),
                        })
                        if len(buffer_images) == 4:
                            flush_grid()
            flush_grid()
            return

        # Fallback text-only
        # Add section title
        section_title = self._extract_section_title(section_data)
        title_para = doc.add_paragraph(section_title)
        title_para.style = doc.styles['Custom Header']

        # Optional images for disclosure page: logo and signature
        if section_data.get('type') == 'disclosure_page':
            from pathlib import Path as _Path
            # Use logo path from disclosure metadata (same directive as cover), fallback to company logo
            logo_path = (section_data.get('metadata', {}) or {}).get('logo_path') or (self.company_info or {}).get('logo_path')
            if logo_path and _Path(logo_path).exists():
                try:
                    doc.add_picture(logo_path, width=Inches(2.0))
                    doc.add_paragraph()
                except Exception:
                    pass
            sig_path = (self.investigator_info or {}).get('signature_path')
            if sig_path and _Path(sig_path).exists():
                try:
                    doc.add_picture(sig_path, width=Inches(2.0))
                    doc.add_paragraph()
                except Exception:
                    pass

        content = section_data.get('content', '')
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
    
    def _extract_section_title(self, section: Dict[str, Any]) -> str:
        """Extract title from section data"""
        
        # Try multiple ways to get the title
        if 'title' in section:
            return section['title']
        
        if 'section_name' in section:
            return section['section_name']
        
        if 'section_id' in section:
            section_id = section['section_id']
            return section_id.replace('_', ' ').title()
        
        # Extract from content
        content = section.get('content', '')
        lines = content.split('\n')
        for line in lines:
            if line.strip() and ('SECTION' in line.upper() or '=' in line):
                return line.strip().rstrip('=').strip()
        
        return "Unknown Section"
    
    def _extract_case_id(self, section_data: Dict[str, Any]) -> str:
        """Extract case ID from section data"""
        
        # Look through section data for case ID
        for section in section_data.values():
            if isinstance(section, dict):
                metadata = section.get('metadata', {})
                if 'case_id' in metadata:
                    return metadata['case_id']
                
                # Check in case_data
                case_data = section.get('case_data', {})
                if 'case_id' in case_data:
                    return case_data['case_id']
        
        # Generate default case ID
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"CASE_{timestamp}"
    
    def _extract_client_name(self, section_data: Dict[str, Any]) -> str:
        """Extract client name from section data"""
        
        # Look through section data for client name
        for section in section_data.values():
            if isinstance(section, dict):
                case_data = section.get('case_data', {})
                if 'client_name' in case_data:
                    return case_data['client_name']
                
                # Check in content
                content = section.get('content', '')
                if 'Client:' in content or 'Client Name:' in content:
                    lines = content.split('\n')
                    for line in lines:
                        if 'Client' in line and ':' in line:
                            return line.split(':')[1].strip()
        
        return "Unknown Client"
    
    def _find_section_cp(self, section_data: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """Find Section CP data from section_data"""
        for v in section_data.values():
            if isinstance(v, dict) and v.get('section_id') == 'section_cp':
                return v
            if isinstance(v, dict) and 'render_data' in v and isinstance(v['render_data'], dict):
                if v['render_data'].get('manifest', {}).get('section_key') == 'section_cp':
                    return v
        return None
    
    def save_report_data(self, report_data: Dict[str, Any], filename: str):
        """Save report data as JSON for later processing"""
        
        with open(filename, 'w') as f:
            json.dump(report_data, f, indent=2, default=str)
        
        logger.info(f"Report data saved to {filename}")
    
    def load_report_data(self, filename: str) -> Dict[str, Any]:
        """Load report data from JSON file"""
        
        with open(filename, 'r') as f:
            report_data = json.load(f)
        
        logger.info(f"Report data loaded from {filename}")
        
        return report_data
