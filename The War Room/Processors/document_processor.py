#!/usr/bin/env python3
"""
Document Processor - OCR, PDF, Image, Video, and Form Processing
Handles intake and processing of all document types for the DKI Engine
"""

import os
import sys
import json
import logging
import hashlib
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import mimetypes

# Lazy loading for heavy dependencies
def _load_cv2():
    """Lazy load OpenCV to reduce import time"""
    try:
        import cv2
        return cv2
    except ImportError:
        return None

def _load_numpy():
    """Lazy load NumPy to reduce import time"""
    try:
        import numpy as np
        return np
    except ImportError:
        return None

# OCR and PDF processing - Multiple OCR Engine Support (Lazy Loading)
PIL_AVAILABLE = False
HAVE_TESSERACT = False
HAVE_EASYOCR = False
HAVE_PADDLEOCR = False

# Lazy loading variables
_pytesseract = None
_easyocr = None
_paddleocr = None
_Image = None
_ImageEnhance = None
_ImageFilter = None

def _load_ocr_modules():
    """Lazy load OCR modules to reduce import time"""
    global PIL_AVAILABLE, HAVE_TESSERACT, HAVE_EASYOCR, HAVE_PADDLEOCR
    global _pytesseract, _easyocr, _paddleocr, _Image, _ImageEnhance, _ImageFilter
    
    # OCR and PDF processing - Multiple OCR Engine Support
    if not PIL_AVAILABLE:
        try:
            import pytesseract
            from PIL import Image, ImageEnhance, ImageFilter
            _pytesseract = pytesseract
            _Image = Image
            _ImageEnhance = ImageEnhance
            _ImageFilter = ImageFilter
            PIL_AVAILABLE = True
            HAVE_TESSERACT = True
        except ImportError:
            pass

    # Alternative OCR: EasyOCR
    if not HAVE_EASYOCR:
        try:
            import easyocr
            _easyocr = easyocr
            HAVE_EASYOCR = True
        except ImportError:
            pass

    # Alternative OCR: PaddleOCR
    if not HAVE_PADDLEOCR:
        try:
            import paddleocr
            _paddleocr = paddleocr
            HAVE_PADDLEOCR = True
        except ImportError:
            pass

# Cloud OCR: Azure Computer Vision
try:
    from azure.cognitiveservices.vision.computervision import ComputerVisionClient
    from azure.cognitiveservices.vision.computervision.models import OperationStatusCodes
    from msrest.authentication import CognitiveServicesCredentials
    HAVE_AZURE_OCR = True
except ImportError:
    HAVE_AZURE_OCR = False

# Determine overall OCR availability
HAVE_OCR = HAVE_TESSERACT or HAVE_EASYOCR or HAVE_PADDLEOCR or HAVE_AZURE_OCR

try:
    import PyPDF2
    import pdfplumber
    HAVE_PDF = True
except ImportError:
    HAVE_PDF = False

# Video processing (lazy loaded)
HAVE_VIDEO = False

# Document processing
try:
    import docx
    from openpyxl import load_workbook
    HAVE_DOCS = True
except ImportError:
    HAVE_DOCS = False

# EXIF and metadata
try:
    from PIL.ExifTags import TAGS
    import piexif
    HAVE_EXIF = True
except ImportError:
    HAVE_EXIF = False

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Comprehensive document processing engine with OCR, metadata extraction, and AI integration"""
    
    # Class-level OCR reader cache for reuse across instances
    _shared_easyocr_reader = None
    _shared_paddle_ocr = None
    _ocr_cache_lock = None
    
    def __init__(self):
        # Initialize thread lock for OCR cache if not already done
        if DocumentProcessor._ocr_cache_lock is None:
            import threading
            DocumentProcessor._ocr_cache_lock = threading.Lock()
        self.supported_formats = {
            'pdf': ['.pdf'],
            'image': ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.heic', '.heif'],
            'video': ['.mp4', '.mov', '.avi', '.mkv', '.wmv', '.flv'],
            'audio': ['.mp3', '.wav', '.m4a', '.aac', '.ogg', '.wma', '.flac'],
            'document': ['.docx', '.doc', '.txt', '.rtf'],
            'spreadsheet': ['.xlsx', '.xls', '.csv'],
            'contract': ['.pdf', '.docx', '.doc'],
            'form': ['.pdf', '.docx', '.doc', '.xlsx', '.xls']
        }
        
        self.ocr_config = {
            'language': 'eng',
            'oem': 3,  # Default OCR Engine Mode
            'psm': 6,  # Page Segmentation Mode - uniform block of text
            'dpi': 300,
            'preprocessing': True
        }
        
        self.processing_stats = {
            'total_files': 0,
            'successful': 0,
            'failed': 0,
            'ocr_processed': 0,
            'metadata_extracted': 0
        }
        
        logger.info("Document Processor initialized")
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check and log available dependencies"""
        deps = {
            'OCR (Tesseract)': HAVE_OCR,
            'PDF Processing': HAVE_PDF,
            'Video Processing': HAVE_VIDEO,
            'Document Processing': HAVE_DOCS,
            'EXIF Metadata': HAVE_EXIF
        }
        
        for dep, available in deps.items():
            status = "Available" if available else "Missing"
            logger.info(f"{dep}: {status}")
    
    def process_files(self, file_list: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Process a list of uploaded files and extract all relevant information"""
        
        logger.info(f"Starting processing of {len(file_list)} files")
        self.processing_stats['total_files'] = len(file_list)
        
        processed_data = {
            'files': {},
            'extracted_text': {},
            'metadata': {},
            'images': {},
            'videos': {},
            'audio': {},
            'contracts': {},
            'forms': {},
            'processing_log': [],
            'summary': {}
        }
        
        for file_info in file_list:
            try:
                result = self._process_single_file(file_info)
                file_id = self._generate_file_id(file_info)
                
                processed_data['files'][file_id] = result
                
                # Categorize by type
                file_type = file_info.get('type', 'unknown')
                if file_type == 'contract':
                    processed_data['contracts'][file_id] = result
                elif file_type == 'form':
                    processed_data['forms'][file_id] = result
                elif file_type == 'image':
                    processed_data['images'][file_id] = result
                elif file_type == 'video':
                    processed_data['videos'][file_id] = result
                elif file_type == 'audio':
                    processed_data['audio'][file_id] = result
                
                # Store extracted text
                if result.get('text'):
                    processed_data['extracted_text'][file_id] = result['text']
                
                # Store metadata
                if result.get('metadata'):
                    processed_data['metadata'][file_id] = result['metadata']
                
                processed_data['processing_log'].append({
                    'file': file_info['name'],
                    'status': 'success',
                    'timestamp': datetime.now().isoformat()
                })
                
                self.processing_stats['successful'] += 1
                
            except Exception as e:
                logger.error(f"Failed to process {file_info['name']}: {str(e)}")
                
                processed_data['processing_log'].append({
                    'file': file_info['name'],
                    'status': 'failed',
                    'error': str(e),
                    'timestamp': datetime.now().isoformat()
                })
                
                self.processing_stats['failed'] += 1
        
        # Generate summary
        processed_data['summary'] = self._generate_processing_summary(processed_data)

        # Aggregate convenience fields for external harnesses
        try:
            total_text_len = sum(len(t or '') for t in processed_data.get('extracted_text', {}).values())
            processed_data['text'] = '\n\n'.join(processed_data.get('extracted_text', {}).values())
            processed_data['success'] = total_text_len > 0
            # Flatten processing methods from first file (best-effort)
            first_file = next(iter(processed_data.get('files', {}).values()), {})
            processed_data['processing_methods'] = first_file.get('processing_methods', []) or []
        except Exception:
            processed_data['success'] = False

        logger.info(f"Processing complete: {self.processing_stats['successful']} successful, {self.processing_stats['failed']} failed")

        return processed_data
    
    def _process_single_file(self, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """Process a single file based on its type and format"""
        
        file_path = file_info['path']
        file_name = file_info['name']
        file_type = file_info.get('type', 'unknown')
        
        logger.debug(f"Processing {file_name} as {file_type}")
        
        result = {
            'file_info': file_info,
            'processing_timestamp': datetime.now().isoformat(),
            'file_hash': self._calculate_file_hash(file_path),
            'mime_type': mimetypes.guess_type(file_path)[0],
            'text': None,
            'metadata': {},
            'images': [],
            'processing_methods': []
        }
        
        # Determine processing method based on file extension
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext in self.supported_formats['pdf']:
            result.update(self._process_pdf(file_path))
        elif file_ext in self.supported_formats['image']:
            result.update(self._process_image(file_path))
        elif file_ext in self.supported_formats['video']:
            result.update(self._process_video(file_path))
        elif file_ext in self.supported_formats['audio']:
            result.update(self._process_audio_placeholder(file_path))
        elif file_ext in self.supported_formats['document']:
            result.update(self._process_document(file_path))
        elif file_ext in self.supported_formats['spreadsheet']:
            result.update(self._process_spreadsheet(file_path))
        else:
            # Try to process as text file
            result.update(self._process_text_file(file_path))
        
        # Extract metadata for all file types
        result['metadata'].update(self._extract_file_metadata(file_path))
        
        return result
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF files - extract text, images, and metadata"""
        
        result = {
            'text': '',
            'images': [],
            'pages': 0,
            'processing_methods': ['pdf_text_extraction']
        }
        
        if not HAVE_PDF:
            logger.warning("PDF processing libraries not available")
            return result
        
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                result['pages'] = len(pdf.pages)
                text_content = []
                
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    
                    # Extract images from page if available
                    if hasattr(page, 'images'):
                        for img in page.images:
                            result['images'].append({
                                'page': page_num + 1,
                                'bbox': img.get('bbox'),
                                'width': img.get('width'),
                                'height': img.get('height')
                            })
                
                result['text'] = '\n\n'.join(text_content)
                result['processing_methods'].append('pdfplumber')
        
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {str(e)}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    result['pages'] = len(pdf_reader.pages)
                    
                    text_content = []
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    
                    result['text'] = '\n\n'.join(text_content)
                    result['processing_methods'].append('PyPDF2')
            
            except Exception as e2:
                logger.error(f"All PDF processing methods failed for {file_path}: {str(e2)}")
                result['processing_methods'].append('failed')
        
        # If no text extracted, try OCR on PDF pages
        if not result['text'].strip() and HAVE_OCR:
            logger.info(f"Attempting OCR on PDF: {file_path}")
            result.update(self._ocr_pdf(file_path))
        
        return result
    
    def _process_image(self, file_path: str) -> Dict[str, Any]:
        """Process image files - OCR and metadata extraction"""
        
        result = {
            'text': '',
            'dimensions': None,
            'processing_methods': ['image_metadata']
        }
        
        try:
            # Load image
            _load_ocr_modules()
            with _Image.open(file_path) as img:
                result['dimensions'] = img.size
                result['mode'] = img.mode
                result['format'] = img.format
                
                # Extract EXIF data
                if HAVE_EXIF and hasattr(img, '_getexif'):
                    exif_data = img._getexif()
                    if exif_data:
                        result['exif'] = {TAGS.get(k, k): v for k, v in exif_data.items()}
                
                # Perform OCR if available
                if HAVE_OCR:
                    result.update(self._perform_ocr(img))
                    result['processing_methods'].append('tesseract_ocr')
        
        except Exception as e:
            logger.error(f"Failed to process image {file_path}: {str(e)}")
            result['processing_methods'].append('failed')
        
        return result
    
    def _process_video(self, file_path: str) -> Dict[str, Any]:
        """Process video files - extract frames, metadata, and OCR on frames"""
        
        result = {
            'text': '',
            'frames_extracted': 0,
            'duration': 0,
            'fps': 0,
            'resolution': None,
            'processing_methods': ['video_metadata']
        }
        
        if not HAVE_VIDEO:
            logger.warning("Video processing libraries not available")
            return result
        
        try:
            cv2 = _load_cv2()
            if cv2 is None:
                return None
            cap = cv2.VideoCapture(file_path)
            
            # Get video properties
            result['fps'] = cap.get(cv2.CAP_PROP_FPS)
            result['duration'] = cap.get(cv2.CAP_PROP_FRAME_COUNT) / result['fps'] if result['fps'] > 0 else 0
            result['resolution'] = (
                int(cap.get(cv2.CAP_PROP_FRAME_WIDTH)),
                int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            )
            
            # Extract frames for OCR (sample every 30 seconds)
            if HAVE_OCR and result['fps'] > 0:
                frame_interval = int(result['fps'] * 30)  # Every 30 seconds
                frame_count = 0
                extracted_text = []
                
                while True:
                    ret, frame = cap.read()
                    if not ret:
                        break
                    
                    if frame_count % frame_interval == 0:
                        # Convert frame to PIL Image and perform OCR
                        frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                        _load_ocr_modules()
                        pil_image = _Image.fromarray(frame_rgb)
                        
                        ocr_result = self._perform_ocr(pil_image)
                        if ocr_result.get('text', '').strip():
                            timestamp = frame_count / result['fps']
                            extracted_text.append(f"[{timestamp:.1f}s] {ocr_result['text']}")
                        
                        result['frames_extracted'] += 1
                    
                    frame_count += 1
                
                result['text'] = '\n\n'.join(extracted_text)
                result['processing_methods'].append('frame_ocr')
            
            cap.release()
        
        except Exception as e:
            logger.error(f"Failed to process video {file_path}: {str(e)}")
            result['processing_methods'].append('failed')
        
        return result

    def _process_audio_placeholder(self, file_path: str) -> Dict[str, Any]:
        """Placeholder entry for audio files; detailed analysis handled by media processing engine"""
        return {
            'processing_methods': ['audio_placeholder'],
            'notes': 'Audio transcription and analysis handled by media processing engine'
        }

    def _process_document(self, file_path: str) -> Dict[str, Any]:
        """Process Word documents and other text documents"""
        
        result = {
            'text': '',
            'processing_methods': []
        }
        
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.docx' and HAVE_DOCS:
                doc = docx.Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                result['text'] = '\n\n'.join(paragraphs)
                result['processing_methods'].append('python-docx')
                
                # Extract tables if present
                if doc.tables:
                    table_text = []
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                            if row_text.strip():
                                table_text.append(row_text)
                    
                    if table_text:
                        result['text'] += '\n\nTables:\n' + '\n'.join(table_text)
            
            elif file_ext in ['.txt', '.rtf']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    result['text'] = f.read()
                result['processing_methods'].append('text_file')
        
        except Exception as e:
            logger.error(f"Failed to process document {file_path}: {str(e)}")
            result['processing_methods'].append('failed')
        
        return result
    
    def _process_spreadsheet(self, file_path: str) -> Dict[str, Any]:
        """Process Excel and CSV files"""
        
        result = {
            'text': '',
            'sheets': [],
            'processing_methods': []
        }
        
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext in ['.xlsx', '.xls'] and HAVE_DOCS:
                wb = load_workbook(file_path, data_only=True)
                sheet_data = []
                
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    sheet_text = f"Sheet: {sheet_name}\n"
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                        if row_text.strip():
                            sheet_text += row_text + '\n'
                    
                    sheet_data.append(sheet_text)
                    result['sheets'].append(sheet_name)
                
                result['text'] = '\n\n'.join(sheet_data)
                result['processing_methods'].append('openpyxl')
            
            elif file_ext == '.csv':
                import csv
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    csv_reader = csv.reader(f)
                    rows = []
                    for row in csv_reader:
                        rows.append(' | '.join(row))
                    result['text'] = '\n'.join(rows)
                result['processing_methods'].append('csv_reader')
        
        except Exception as e:
            logger.error(f"Failed to process spreadsheet {file_path}: {str(e)}")
            result['processing_methods'].append('failed')
        
        return result
    
    def _process_text_file(self, file_path: str) -> Dict[str, Any]:
        """Process plain text files"""
        
        result = {
            'text': '',
            'processing_methods': ['text_file']
        }
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        result['text'] = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not result['text']:
                # Final fallback - read as binary and decode with errors ignored
                with open(file_path, 'rb') as f:
                    result['text'] = f.read().decode('utf-8', errors='ignore')
        
        except Exception as e:
            logger.error(f"Failed to process text file {file_path}: {str(e)}")
            result['processing_methods'].append('failed')
        
        return result
    
    def _perform_ocr(self, image) -> Dict[str, Any]:
        """Perform OCR using multiple engines with fallback support"""
        
        result = {
            'text': '',
            'confidence': 0,
            'processing_methods': [],
            'engine_used': None,
            'fallback_attempts': []
        }
        
        if not HAVE_OCR:
            logger.warning("No OCR engines available")
            return result
        
        # Try OCR engines in order of preference
        ocr_engines = [
            ('easyocr', self._ocr_with_easyocr),
            ('paddleocr', self._ocr_with_paddleocr),
            ('tesseract', self._ocr_with_tesseract),
            ('azure', self._ocr_with_azure)
        ]
        
        for engine_name, engine_func in ocr_engines:
            try:
                engine_result = engine_func(image)
                if engine_result and engine_result.get('text', '').strip():
                    result.update(engine_result)
                    result['engine_used'] = engine_name
                    result['processing_methods'].append(f'ocr_{engine_name}')
                    self.processing_stats['ocr_processed'] += 1
                    logger.info(f"OCR successful with {engine_name} engine")
                    return result
                else:
                    result['fallback_attempts'].append(f'{engine_name}_no_text')
            except Exception as e:
                logger.warning(f"OCR engine {engine_name} failed: {str(e)}")
                result['fallback_attempts'].append(f'{engine_name}_error')
                continue
        
        # If all engines failed
        result['processing_methods'].append('ocr_all_failed')
        logger.error("All OCR engines failed to extract text")
        return result
    
    def _preprocess_image_for_ocr(self, image):
        """Preprocess image to improve OCR accuracy"""
        
        try:
            # Convert to grayscale
            if image.mode != 'L':
                image = image.convert('L')
            
            # Enhance contrast
            _load_ocr_modules()
            enhancer = _ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)
            
            # Sharpen
            image = image.filter(_ImageFilter.SHARPEN)
            
            # Resize if too small (OCR works better on larger images)
            width, height = image.size
            if width < 300 or height < 300:
                scale_factor = max(300 / width, 300 / height)
                new_size = (int(width * scale_factor), int(height * scale_factor))
                image = image.resize(new_size, _Image.Resampling.LANCZOS)
        
        except Exception as e:
            logger.warning(f"Image preprocessing failed: {str(e)}")
        
        return image
    
    def _ocr_with_easyocr(self, image) -> Optional[Dict[str, Any]]:
        """Perform OCR using EasyOCR engine"""
        if not HAVE_EASYOCR:
            return None
        
        try:
            np = _load_numpy()
            if np is None:
                return None
            # Use shared EasyOCR reader (cached across instances)
            with DocumentProcessor._ocr_cache_lock:
                if DocumentProcessor._shared_easyocr_reader is None:
                    _load_ocr_modules()
                    DocumentProcessor._shared_easyocr_reader = _easyocr.Reader(['en', 'es'], gpu=False)
                reader = DocumentProcessor._shared_easyocr_reader
            
            # Convert PIL image to numpy array
            img_array = np.array(image)
            
            # Perform OCR
            results = reader.readtext(img_array)
            
            # Extract text and confidence
            text_parts = []
            confidences = []
            
            for (bbox, text, confidence) in results:
                if confidence > 0.5:  # Filter low confidence results
                    text_parts.append(text)
                    confidences.append(confidence * 100)  # Convert to percentage
            
            return {
                'text': ' '.join(text_parts),
                'confidence': sum(confidences) / len(confidences) if confidences else 0,
                'engine': 'easyocr',
                'details': {'results_count': len(results), 'filtered_count': len(text_parts)}
            }
            
        except Exception as e:
            logger.error(f"EasyOCR failed: {e}")
            return None
    
    def _ocr_with_paddleocr(self, image) -> Optional[Dict[str, Any]]:
        """Perform OCR using PaddleOCR engine"""
        if not HAVE_PADDLEOCR:
            return None
        
        try:
            np = _load_numpy()
            if np is None:
                return None
            # Use shared PaddleOCR (cached across instances)
            with DocumentProcessor._ocr_cache_lock:
                if DocumentProcessor._shared_paddle_ocr is None:
                    _load_ocr_modules()
                    DocumentProcessor._shared_paddle_ocr = _paddleocr.PaddleOCR(use_angle_cls=True, lang='en', use_gpu=False)
                paddle_reader = DocumentProcessor._shared_paddle_ocr
            
            # Convert PIL image to numpy array
            img_array = np.array(image)
            
            # Perform OCR
            results = paddle_reader.ocr(img_array, cls=True)
            
            # Extract text and confidence
            text_parts = []
            confidences = []
            
            if results and results[0]:
                for line in results[0]:
                    if line and len(line) >= 2:
                        text = line[1][0] if line[1] else ''
                        confidence = line[1][1] if line[1] and len(line[1]) > 1 else 0
                        
                        if confidence > 0.5 and text.strip():
                            text_parts.append(text)
                            confidences.append(confidence * 100)
            
            return {
                'text': ' '.join(text_parts),
                'confidence': sum(confidences) / len(confidences) if confidences else 0,
                'engine': 'paddleocr',
                'details': {'lines_processed': len(results[0]) if results and results[0] else 0}
            }
            
        except Exception as e:
            logger.error(f"PaddleOCR failed: {e}")
            return None
    
    def _ocr_with_tesseract(self, image) -> Optional[Dict[str, Any]]:
        """Perform OCR using Tesseract engine"""
        if not HAVE_TESSERACT:
            return None
        
        try:
            # Configure Tesseract path for Windows
            if sys.platform.startswith('win'):
                possible_paths = [
                    r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                    r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe'
                ]
                for path in possible_paths:
                    if Path(path).exists():
                        _load_ocr_modules()
                        _pytesseract.pytesseract.tesseract_cmd = path
                        break
            
            # Preprocess image for better OCR
            if self.ocr_config['preprocessing']:
                image = self._preprocess_image_for_ocr(image)
            
            # Configure tesseract
            config = f"--oem {self.ocr_config['oem']} --psm {self.ocr_config['psm']}"
            
            # Perform OCR
            _load_ocr_modules()
            text = _pytesseract.image_to_string(image, lang=self.ocr_config['language'], config=config)
            
            # Get confidence data if available
            confidence = 0
            try:
                data = _pytesseract.image_to_data(image, output_type=_pytesseract.Output.DICT)
                confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
                if confidences:
                    confidence = sum(confidences) / len(confidences)
            except:
                confidence = 85  # Default confidence for Tesseract
            
            return {
                'text': text.strip(),
                'confidence': confidence,
                'engine': 'tesseract',
                'details': {'config_used': config}
            }
            
        except Exception as e:
            logger.error(f"Tesseract OCR failed: {e}")
            return None
    
    def _ocr_with_azure(self, image) -> Optional[Dict[str, Any]]:
        """Perform OCR using Azure Computer Vision API"""
        if not HAVE_AZURE_OCR:
            return None
        
        try:
            # This would require Azure credentials
            # For now, return None - can be implemented when Azure keys are available
            logger.info("Azure OCR available but not configured")
            return None
            
        except Exception as e:
            logger.error(f"Azure OCR failed: {e}")
            return None
    
    def _ocr_pdf(self, file_path: str) -> Dict[str, Any]:
        """Perform OCR on PDF pages"""
        
        result = {
            'text': '',
            'processing_methods': ['pdf_ocr']
        }
        
        # This would require converting PDF pages to images first
        # Implementation depends on additional libraries like pdf2image
        logger.info(f"PDF OCR not fully implemented for {file_path}")
        
        return result
    
    def _extract_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract file system and format-specific metadata"""
        
        metadata = {}
        
        try:
            stat = os.stat(file_path)
            metadata.update({
                'file_size': stat.st_size,
                'created_time': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified_time': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'accessed_time': datetime.fromtimestamp(stat.st_atime).isoformat()
            })
            
            self.processing_stats['metadata_extracted'] += 1
        
        except Exception as e:
            logger.warning(f"Failed to extract metadata for {file_path}: {str(e)}")
        
        return metadata
    
    def _calculate_file_hash(self, file_path: str) -> Dict[str, str]:
        """Calculate MD5 and SHA256 hashes for file integrity"""
        
        hashes = {}
        
        try:
            with open(file_path, 'rb') as f:
                data = f.read()
                hashes['md5'] = hashlib.md5(data).hexdigest()
                hashes['sha256'] = hashlib.sha256(data).hexdigest()
        
        except Exception as e:
            logger.error(f"Failed to calculate hash for {file_path}: {str(e)}")
        
        return hashes
    
    def _generate_file_id(self, file_info: Dict[str, Any]) -> str:
        """Generate a unique ID for the file"""
        
        name = file_info['name']
        path = file_info['path']
        timestamp = file_info.get('uploaded_date', datetime.now().isoformat())
        
        # Create hash from name + path + timestamp
        id_string = f"{name}_{path}_{timestamp}"
        return hashlib.md5(id_string.encode()).hexdigest()[:12]
    
    def _generate_processing_summary(self, processed_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate a summary of the processing results"""
        
        return {
            'total_files_processed': len(processed_data['files']),
            'successful_extractions': len([f for f in processed_data['files'].values() 
                                          if f.get('text', '').strip()]),
            'files_with_metadata': len(processed_data['metadata']),
            'contracts_processed': len(processed_data['contracts']),
            'forms_processed': len(processed_data['forms']),
            'images_processed': len(processed_data['images']),
            'videos_processed': len(processed_data['videos']),
            'audio_processed': len(processed_data.get('audio', {})),
            'total_text_length': sum(len(text) for text in processed_data['extracted_text'].values()),
            'processing_stats': self.processing_stats.copy(),
            'timestamp': datetime.now().isoformat()
        }
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get current processing statistics"""
        return self.processing_stats.copy()
    
    def configure_ocr(self, **kwargs):
        """Configure OCR settings"""
        for key, value in kwargs.items():
            if key in self.ocr_config:
                self.ocr_config[key] = value
                logger.info(f"OCR config updated: {key} = {value}")
    
    def is_supported_format(self, file_path: str) -> bool:
        """Check if file format is supported"""
        ext = Path(file_path).suffix.lower()
        return any(ext in formats for formats in self.supported_formats.values())
